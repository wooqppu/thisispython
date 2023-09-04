
from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

load_wb = load_workbook("수료증명단.xlsx")
load_ws = load_wb.active

name_list=[]
birthday_list=[]
ho_list = []

for i in range(1, load_ws.max_row+1) :
    name_list.append(load_ws.cell(i, 1).value)
    birthday_list.append(load_ws.cell(i, 2).value)
    ho_list.append(load_ws.cell(i, 3).value)

print(name_list)
print(birthday_list)
print(ho_list)

for i in range(len(name_list)):

    doc = docx.Document("수료증양식.docx")


    para = doc.add_paragraph()
    run = para.add_run("\n\n")
    run = para.add_run("           제 " + ho_list[i] + "호 \n")
    run.font.size = docx.shared.Pt(15)

    para = doc.add_paragraph()
    run = para.add_run("\n\n")
    run = para.add_run(" 수 료 증 \n")
    run.font.size = docx.shared.Pt(40)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("\n\n")
    run = para.add_run("       성명 : " + name_list[i] + "  \n")
    run.font.size = docx.shared.Pt(20)

    run = para.add_run("        생년월일 : " + birthday_list[i] + "  \n")
    run.font.size = docx.shared.Pt(15)
    run = para.add_run("        교육과정 : 파이썬 프로그래밍")
    run.font.size = docx.shared.Pt(15)

    para = doc.add_paragraph()
    run = para.add_run("\n\n")
    run = para.add_run(" 위 사람은 해당 교육 과정을 성실하게 수료함 \n")
    run.font.size = docx.shared.Pt(23)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


    para = doc.add_paragraph()
    run = para.add_run("\n\n")
    run = para.add_run(" 2023.03.30 \n")
    run.font.size = docx.shared.Pt(20)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph()
    run = para.add_run("\n\n")
    run = para.add_run(" 그린컴퓨터아카데미 울산점 \n")
    run.font.size = docx.shared.Pt(18)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(name_list[i] + ".docx")
    convert(name_list[i] + ".docx", name_list[i] + ".pdf")